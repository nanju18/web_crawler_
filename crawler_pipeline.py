import os
import asyncio
from typing import Optional, List
import datetime
from fastapi import APIRouter, HTTPException
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from pydantic import BaseModel, Field, model_validator
from dotenv import load_dotenv
from best_first_test import BestFirstCrawl
from depth_first import DepthFirstCrawl
from breath_first import BreathFirstCrawl

load_dotenv()

router = APIRouter()
asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

class CrawlRequest(BaseModel):
    url: str
    strategy: str
    method: str
    #keywords: Optional[List[str]] = None
    depth: int = Field(..., ge=0, le=3)
    '''
    @model_validator(mode="before")
    def validate_keywords_for_strategy(cls, values):
        strategy = values.get("strategy", "").lower()
        keywords = values.get("keywords")

        if strategy == "best first" and not keywords:
            raise ValueError("Keywords must be provided for best first strategy.")
        if strategy == "depth first" and keywords:
            print("Note: Keywords provided but will be ignored for depth first strategy.")
        return values
'''
def save_results_to_docx(strategy, method, results: list[dict]) -> str:
    local_storage_path = os.path.join(os.path.expanduser("~"), "Downloads", "crawl_exports")
    os.makedirs(local_storage_path, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%H%M%S")
    file_path = os.path.join(local_storage_path, f"crawl_result_{strategy}_{method}_{timestamp}.docx")

    doc = Document()
    doc.add_heading("Crawled Content", level=1)

    for i, item in enumerate(results, start=1):
        doc.add_heading(f"{i}. {item['url']}", level=2)
        text = item.get("fit_markdown")
        if text:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=5000,
                chunk_overlap=1000,
                separators=["\n\n", "\n", " ", "", "."]
            )
            chunks = text_splitter.split_text(text)
            for chunk in chunks:
                doc.add_paragraph(chunk)
        else:
            doc.add_paragraph("No content available.")

    print(f"Saving DOCX to: {file_path}")
    doc.save(file_path)
    return file_path

@router.post("/")
def start_crawling(request: CrawlRequest):
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        url = request.url
        method = request.method
        strategy = request.strategy.lower()
        depth = request.depth
        #keywords = request.keywords
        results = []

        if strategy == "depth first":
            if method == "single":
                crawl_single_page_service = DepthFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)

            elif method == "recursive":
                depth_first_crawl_service = DepthFirstCrawl()
                depth_first_crawl = depth_first_crawl_service.depth_first_crawl(url, depth)
                results = loop.run_until_complete(depth_first_crawl)

        elif strategy == "breath first":
            if method == "single":
                crawl_single_page_service = BreathFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)

            elif method == "recursive":
                depth_first_crawl_service = BreathFirstCrawl()
                depth_first_crawl = depth_first_crawl_service.Breath_first_crawl(url, depth)
                results = loop.run_until_complete(depth_first_crawl)

        else:
            raise HTTPException(status_code=400, detail="Invalid strategy")

        save_results_to_docx(strategy, method, results)
        return {
            "message": "Crawling completed and data stored successfully.",
            "strategy": strategy,
            "method": method,
            "pages_crawled": len(results or [])
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
